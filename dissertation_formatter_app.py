# dissertation_formatter_app.py (ECU Custom)
import streamlit as st
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import tempfile
import os

ECU_PURPLE = "#4B1869"
ECU_GOLD = "#FDB913"
ECU_LOGO_URL = "https://brand.ecu.edu/wp-content/uploads/sites/54/2021/05/ECUlogo_horizontal_purplegold.png"

def check_font(document):
    issues = []
    for para in document.paragraphs:
        for run in para.runs:
            if run.font.name not in ["Times New Roman", "Arial"]:
                issues.append((para.text[:50], run.font.name))
    return issues

def check_spacing(document):
    issues = []
    for para in document.paragraphs:
        if para.paragraph_format.space_before != Pt(0) or para.paragraph_format.space_after != Pt(0):
            issues.append(para.text[:50])
    return issues

def check_alignment(document):
    misaligned = []
    for para in document.paragraphs:
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            misaligned.append(para.text[:50])
    return misaligned

def check_heading_levels(document):
    heading_issues = []
    headings_found = set()
    for para in document.paragraphs:
        if para.style.name.startswith("Heading"):
            headings_found.add(para.style.name)
    for level in ["Heading 1", "Heading 2"]:
        if level not in headings_found:
            heading_issues.append(f"Missing required heading level: {level}")
    return heading_issues

def check_appendices(document):
    has_irb = any("appendix a" in para.text.lower() and "irb" in para.text.lower() for para in document.paragraphs)
    return [] if has_irb else ["Missing Appendix A: IRB approval letter"]

def add_page_numbers(document):
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def insert_table_of_contents(document):
    toc_paragraph = document.paragraphs[0].insert_paragraph_before()
    run = toc_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def format_and_review(uploaded_file):
    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, uploaded_file.name)
    with open(docx_path, 'wb') as f:
        f.write(uploaded_file.read())
    doc = docx.Document(docx_path)
    feedback = {
        "Font Issues": check_font(doc),
        "Spacing Issues": check_spacing(doc),
        "Alignment Issues": check_alignment(doc),
        "Heading Level Issues": check_heading_levels(doc),
        "Appendix Check": check_appendices(doc)
    }
    add_page_numbers(doc)
    insert_table_of_contents(doc)
    output_path = os.path.join(temp_dir, "Formatted_Dissertation.docx")
    doc.save(output_path)
    return feedback, output_path

st.markdown(f"""
    <div style='text-align: center;'>
        <img src='{ECU_LOGO_URL}' width='400'>
        <h1 style='color: {ECU_PURPLE};'>ECU Dissertation Formatter Bot</h1>
        <p style='font-size:18px;'>Helping doctoral students format their dissertation in practice with precision and Pirate pride.</p>
    </div>
    <hr style='border-top: 3px solid {ECU_GOLD}; margin-top: 1em; margin-bottom: 1em;'>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Upload your .docx dissertation file", type=["docx"])

if uploaded_file:
    feedback, output_path = format_and_review(uploaded_file)
    st.success("✅ Formatting completed. Download your revised document below.")
    with open(output_path, "rb") as f:
        st.download_button("⬇️ Download Formatted Dissertation", f, file_name="Formatted_Dissertation.docx")

    st.subheader("📝 Formatting Feedback")
    for issue_type, details in feedback.items():
        st.markdown(f"### {issue_type}")
        if details:
            for item in details:
                st.markdown(f"- {item}")
        else:
            st.markdown("- No issues found")
